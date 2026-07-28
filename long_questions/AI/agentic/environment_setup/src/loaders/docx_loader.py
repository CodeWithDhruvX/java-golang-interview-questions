"""
DOCX document loader
"""
from .base_loader import BaseDocumentLoader
from typing import Dict, Any
from docx import Document


class DOCXLoader(BaseDocumentLoader):
    """Load and extract text from DOCX files"""
    
    def __init__(self, file_path: str):
        super().__init__(file_path)
        if self.file_path.suffix.lower() not in ['.docx', '.doc']:
            raise ValueError("File must be a DOCX")
    
    def load(self) -> Dict[str, Any]:
        """Load DOCX with full metadata"""
        metadata = self.get_metadata()
        text = self.extract_text()
        
        # Additional DOCX-specific metadata
        docx_metadata = self._extract_docx_metadata()
        
        return {
            "content": text,
            "metadata": {**metadata, **docx_metadata},
            "paragraph_count": self._get_paragraph_count(),
            "table_count": self._get_table_count()
        }
    
    def extract_text(self) -> str:
        """Extract text from document"""
        doc = Document(str(self.file_path))
        
        text_content = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            text_content.append("\n--- Table ---")
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                text_content.append(row_text)
        
        return "\n\n".join(text_content)
    
    def _extract_docx_metadata(self) -> Dict[str, str]:
        """Extract DOCX-specific metadata"""
        doc = Document(str(self.file_path))
        core_props = doc.core_properties
        
        return {
            "title": core_props.title or 'Unknown',
            "author": core_props.author or 'Unknown',
            "created": str(core_props.created) if core_props.created else 'Unknown',
            "modified": str(core_props.modified) if core_props.modified else 'Unknown',
        }
    
    def _get_paragraph_count(self) -> int:
        """Get total paragraph count"""
        doc = Document(str(self.file_path))
        return len(doc.paragraphs)
    
    def _get_table_count(self) -> int:
        """Get total table count"""
        doc = Document(str(self.file_path))
        return len(doc.tables)
