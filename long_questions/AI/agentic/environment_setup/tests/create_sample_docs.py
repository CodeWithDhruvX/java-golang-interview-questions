"""
Create sample documents for testing
"""
from docx import Document
import pandas as pd
from pathlib import Path


def create_sample_docs():
    """Create sample documents for testing"""
    sample_dir = Path("sample_docs")
    sample_dir.mkdir(exist_ok=True)
    
    # Create sample TXT
    with open(sample_dir / "sample.txt", "w", encoding="utf-8") as f:
        f.write("This is a sample text file.\nIt contains multiple lines.\nUsed for testing document loading.")
    
    # Create sample CSV
    df = pd.DataFrame({
        "Name": ["Alice", "Bob", "Charlie"],
        "Age": [25, 30, 35],
        "City": ["New York", "London", "Paris"]
    })
    df.to_csv(sample_dir / "sample.csv", index=False)
    
    # Create sample DOCX
    doc = Document()
    doc.add_heading("Sample Document", 0)
    doc.add_paragraph("This is a sample paragraph.")
    doc.add_paragraph("Another paragraph with different content.")
    
    table = doc.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "Header 1"
    table.rows[0].cells[1].text = "Header 2"
    table.rows[1].cells[0].text = "Data 1"
    table.rows[1].cells[1].text = "Data 2"
    table.rows[2].cells[0].text = "Data 3"
    table.rows[2].cells[1].text = "Data 4"
    
    doc.save(sample_dir / "sample.docx")
    
    print("✅ Sample documents created in sample_docs/")


if __name__ == "__main__":
    create_sample_docs()
